import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Styles globaux
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x41, 0x55)

# Titre Principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("🏥 SantéApp – Système de Gestion Médicale")
run_title.font.size = Pt(22)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Documentation Technique & Functional Spec (PWA & FastAPI)")
run_sub.font.size = Pt(12)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

doc.add_paragraph() # Espacement

def add_heading_1(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
    return h

def add_heading_2(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(text)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    return h

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run_b = p.add_run(bold_prefix + " ")
    run_b.bold = True
    p.add_run(text)

# 1. Fonctionnalités
add_heading_1("🚀 1. Fonctionnalités Principales")

add_heading_2("👤 Espace Patient")
add_bullet("Authentification sécurisée :", "Inscription et connexion avec chiffrement BCrypt des mots de passe.")
add_bullet("Prise de Rendez-vous :", "Sélection d'un praticien par spécialité et envoi du motif de consultation.")
add_bullet("Historique Médical :", "Suivi des consultations passées, diagnostics et ordonnances.")
add_bullet("Impression d'Ordonnances :", "Génération et impression au format PDF directement depuis l'application.")

add_heading_2("🩺 Espace Médecin")
add_bullet("Tableau de Bord Consultation :", "Visualisation dynamique des rendez-vous en attente.")
add_bullet("Saisie du Diagnostic :", "Enregistrement des symptômes, diagnostics et prescriptions.")
add_bullet("Clôture du Rendez-vous :", "Mise à jour automatique de l'état du rendez-vous (EN_ATTENTE ➔ TERMINE).")

add_heading_2("📱 Technologie PWA & Notifications")
add_bullet("Installation PWA :", "Bouton d'installation native sur Android, iOS et Desktop.")
add_bullet("Mode Hors Ligne :", "Mise en cache des ressources statiques via Service Worker (sw.js).")
add_bullet("Interface Dynamique (Toast UI) :", "Notifications animées et élégantes pour le suivi des actions utilisateur.")

# 2. Stack Technique
add_heading_1("🛠️ 2. Stack Technique")
add_bullet("Backend :", "Python 3.10+, FastAPI, Uvicorn, SQLAlchemy")
add_bullet("Base de données :", "PostgreSQL (Production via Supabase / Render) / SQLite (Développement local)")
add_bullet("Migrations :", "Alembic")
add_bullet("Sécurité :", "JWT (JSON Web Tokens), Passlib (BCrypt)")
add_bullet("Frontend :", "HTML5, JavaScript Modern (ES6+), Tailwind CSS (CDN)")
add_bullet("PWA :", "Web App Manifest, Service Worker")

# 3. Structure
add_heading_1("📂 3. Structure du Projet")
struct_text = (
    "sante_backend/\n"
    "├── app/\n"
    "│   ├── main.py              # FastAPI app & points d'entrée API\n"
    "│   ├── database.py          # Configuration SQLAlchemy PostgreSQL\n"
    "│   ├── models.py            # Modèles (User, RDV, Consultation)\n"
    "│   ├── schemas.py           # Validations Pydantic\n"
    "│   ├── auth.py              # Gestion JWT & sécurité\n"
    "│   └── routers/\n"
    "│       ├── auth_router.py   # Inscription & Connexion\n"
    "│       ├── patient.py       # Endpoints espace Patient\n"
    "│       └── medecin.py       # Endpoints espace Médecin\n"
    "├── static/\n"
    "│   ├── index.html           # Interface SPA (Single Page Application)\n"
    "│   ├── manifest.json        # Manifeste PWA\n"
    "│   └── sw.js                # Service Worker PWA\n"
    "├── alembic/                 # Script de migrations DB\n"
    "├── requirements.txt         # Dépendances du projet\n"
    "├── Procfile                 # Script de lancement Render\n"
    "└── README.md                # Documentation GitHub"
)

p_code = doc.add_paragraph()
p_code.paragraph_format.space_before = Pt(4)
p_code.paragraph_format.space_after = Pt(10)
r_code = p_code.add_run(struct_text)
r_code.font.name = 'Courier New'
r_code.font.size = Pt(8.5)

# 4. Installation
add_heading_1("⚙️ 4. Installation et Démarrage Local")
add_bullet("1. Cloner le projet :", "git clone <url-du-repo>")
add_bullet("2. Environnement virtuel :", "python -m venv venv && source venv/bin/activate")
add_bullet("3. Dépendances :", "pip install -r requirements.txt")
add_bullet("4. Migrations DB :", "alembic upgrade head")
add_bullet("5. Lancer le serveur :", "uvicorn app.main:app --reload")

# Sauvegarde
doc.save("README_SanteApp.docx")
print("✅ Document Word créé avec succès : README_SanteApp.docx")